import streamlit as st
import pandas as pd
import math
import os
import plotly.graph_objects as go
import numpy as np
from PIL import Image
import io
from docx import Document
from docx.shared import Pt, RGBColor

st.set_page_config(page_title="Dimensionamento Pluvial - NBR 10844", layout="wide")

# ==============================================================================
# OTIMIZAÇÃO DE PERFORMANCE (CACHE PARA IMAGENS)
# ==============================================================================
@st.cache_data
def carregar_imagem(caminho):
    return Image.open(caminho)

# ==============================================================================
# BARRA LATERAL - AUTOR E AVISO LEGAL
# ==============================================================================
st.sidebar.title("👨‍💻 Desenvolvido por")
st.sidebar.markdown("### **João Luiz**")
st.sidebar.markdown("📧 joaoluiz@outlook.com")
st.sidebar.markdown("---")
st.sidebar.warning("⚠️ **Aviso Legal:** Este software é uma ferramenta auxiliar de cálculo. O usuário é o único e exclusivo responsável pela conferência dos resultados e pela aplicação técnica destes dimensionamentos em seus projetos.")
st.sidebar.info("Memorial de Cálculo Automático estruturado conforme os parâmetros da NBR 10844/1989.")

st.title("Dimensionamento de Instalações Prediais de Águas Pluviais")
st.markdown("**Norma de Referência:** NBR 10844/1989")

tab1, tab2, tab3 = st.tabs([
    "1. Calhas Horizontais", 
    "2. Condutores Verticais", 
    "3. Condutores Horizontais"
])

# ==============================================================================
# FUNÇÕES DE GERAÇÃO DE WORD (.DOCX)
# ==============================================================================
def formatar_formula_word(doc, rotulo, formula):
    p = doc.add_paragraph()
    p.add_run(rotulo).bold = True
    run_form = p.add_run(formula)
    run_form.font.name = 'Courier New'
    run_form.font.size = Pt(11)
    
def criar_cabecalho_word(doc, titulo):
    doc.add_heading(titulo, 0)
    p = doc.add_paragraph()
    p.add_run("Desenvolvido por: ").bold = True
    p.add_run("João Luiz\n")
    p.add_run("E-mail: ").bold = True
    p.add_run("joaoluiz@outlook.com\n")
    p.add_run("Norma: ").bold = True
    p.add_run("NBR 10844/1989\n")
    doc.add_paragraph("--------------------------------------------------")
    
    p_aviso = doc.add_paragraph()
    p_aviso.add_run("AVISO: O usuário é o responsável técnico por conferir, validar e assumir a autoria dos cálculos aplicados através deste memorial.").italic = True
    
    doc.add_paragraph("--------------------------------------------------")

# ==============================================================================
# ABA 1: CALHAS HORIZONTAIS
# ==============================================================================
with tab1:
    st.header("Dimensionamento de Calhas Horizontais")
    col_in, col_out = st.columns([1, 1.2])
    
    with col_in:
        st.subheader("Parâmetros de Entrada")
        st.markdown("**Área de Contribuição**")
        area_c1 = st.number_input("Área de contribuição total (A) [m²]:", min_value=0.0, value=150.0, step=10.0, key="area_calha")
        i_pluv1 = st.number_input("Intensidade pluviométrica (I) [mm/h]:", min_value=0.0, value=150.0, step=10.0, key="i_calha")
        
        st.markdown("---")
        st.markdown("**Geometria e Características da Calha (Valores em mm)**")
        tipo_calha = st.radio("Selecione o formato da calha:", ["Retangular", "Semicircular"])
        
        if tipo_calha == "Retangular":
            b_calha_mm = st.number_input("Largura da base da calha (b) [mm]:", min_value=10.0, value=250.0, step=10.0)
            h_total_calha_mm = st.number_input("Altura total da calha (H) [mm]:", min_value=10.0, value=250.0, step=10.0)
            prop_selecionada = st.selectbox("Proporção para a altura molhada:", ["50%", "66%", "75%"])
            b_calha = b_calha_mm / 1000.0
            h_total_calha = h_total_calha_mm / 1000.0
            mapa_proporcoes = {"50%": 0.50, "66%": 0.66, "75%": 0.75}
            fator_prop = mapa_proporcoes[prop_selecionada]
            h_molhada = h_total_calha * fator_prop
        else:
            d_calha_mm = st.number_input("Diâmetro interno da calha semicircular (D) [mm]:", min_value=10.0, value=150.0, step=10.0)
            d_calha = d_calha_mm / 1000.0
            r_calha = d_calha / 2.0
            h_molhada = (2/3) * r_calha 
            
        n_rug = st.number_input("Coeficiente de rugosidade de Manning (n):", min_value=0.001, value=0.011, step=0.001, format="%.3f")
        declividade = st.number_input("Declividade da calha (i) [m/m]:", min_value=0.0001, value=0.0050, step=0.0005, format="%.4f")

    with col_out:
        st.subheader("📋 MEMORIAL DE CÁLCULO DETALHADO")
        
        st.markdown("### 1. Vazão de Projeto ($Q_{proj}$)")
        q_proj1 = (i_pluv1 * area_c1) / 60.0
        
        with st.container():
            st.markdown("**💡 1. Símbolos e Valores Cadastrados:**")
            st.markdown(f"* $I$ = **{i_pluv1:.2f}** mm/h\n* $A$ = **{area_c1:.2f}** m²")
            st.markdown("**📖 2. Fórmula Base:**")
            st.latex(r"Q_{proj} = \frac{I \cdot A}{60}")
            st.markdown("**🧪 3. Substituição Numérica:**")
            st.latex(f"Q_{{proj}} = \\frac{{ {i_pluv1:.2f} \\cdot {area_c1:.2f} }}{{ 60 }}")
            st.metric(label="Vazão Solicitada Calculada", value=f"{q_proj1:.2f} L/min")
        
        st.markdown("---")
        st.markdown(f"### 2. Capacidade da Calha ({tipo_calha}) ($Q_{{calha}}$)")
        K_cte = 60000.0
        
        if tipo_calha == "Retangular":
            S_area = b_calha * h_molhada
            P_perimetro = b_calha + (2 * h_molhada)
            R_hidr = S_area / P_perimetro if P_perimetro > 0 else 0
            with st.container():
                st.markdown("**💡 1. Símbolos e Parâmetros Geométricos:**")
                st.markdown(f"* $b$ = **{b_calha:.3f} m**\n* $H_{{total}}$ = **{h_total_calha:.3f} m**\n* $h$ = **{h_molhada:.3f} m**\n* $S$ = **{S_area:.4f} m²**\n* $P$ = **{P_perimetro:.4f} m**")
        else:
            cos_metade_theta = (r_calha - h_molhada) / r_calha
            theta = 2 * math.acos(cos_metade_theta)
            S_area = (r_calha**2 / 2) * (theta - math.sin(theta))
            P_perimetro = r_calha * theta
            R_hidr = S_area / P_perimetro if P_perimetro > 0 else 0
            with st.container():
                st.markdown("**💡 1. Símbolos e Parâmetros Geométricos:**")
                st.markdown(f"* $D$ = **{d_calha:.3f} m**\n* $R$ = **{r_calha:.3f} m**\n* $h$ = **{h_molhada:.3f} m**\n* $S$ = **{S_area:.4f} m²**\n* $P$ = **{P_perimetro:.4f} m**")

        with st.container():
            st.markdown("**💧 Variáveis Hidráulicas de Fluxo:**")
            st.markdown(f"* $R_h$ = **{R_hidr:.4f} m**\n* $n$ = **{n_rug:.3f}**\n* $i$ = **{declividade:.4f} m/m**\n* $K$ = **{K_cte:.0f}**")
            
            q_calha_calc = (K_cte * S_area * (R_hidr ** (2/3)) * (declividade ** 0.5)) / n_rug

            st.markdown("**📖 2. Fórmula de Manning-Strickler:**")
            st.latex(r"Q_{calha} = \frac{K \cdot S \cdot R_h^{2/3} \cdot i^{0.5}}{n}")
            st.markdown("**🧪 3. Substituição Numérica Completa:**")
            st.latex(f"Q_{{calha}} = \\frac{{ {K_cte:.0f} \\cdot {S_area:.4f} \\cdot {R_hidr:.4f}^{{2/3}} \\cdot {declividade:.4f}^{{0.5}} }}{{ {n_rug:.3f} }}")
            st.metric(label="Capacidade de Escoamento (Q_calha)", value=f"{q_calha_calc:.2f} L/min")
        
        st.markdown("---")
        if q_calha_calc >= q_proj1:
            st.success(f"✅ **CONFORMIDADE ATENDIDA:** A capacidade da calha ({q_calha_calc:.2f} L/min) suporta a vazão ({q_proj1:.2f} L/min).")
        else:
            st.error(f"❌ **NÃO ATENDE:** A capacidade da calha ({q_calha_calc:.2f} L/min) é menor que a vazão de projeto ({q_proj1:.2f} L/min).")

        # GERAÇÃO DO ARQUIVO WORD
        doc_calha = Document()
        criar_cabecalho_word(doc_calha, 'Memorial de Cálculo - Calhas Horizontais')
        doc_calha.add_heading('1. VAZÃO DE PROJETO (Q_proj)', level=1)
        doc_calha.add_paragraph(f"Intensidade (I) = {i_pluv1:.2f} mm/h\nÁrea (A) = {area_c1:.2f} m²")
        formatar_formula_word(doc_calha, "Fórmula Base: ", "Q_proj = (I * A) / 60")
        formatar_formula_word(doc_calha, "Substituição: ", f"Q_proj = ({i_pluv1:.2f} * {area_c1:.2f}) / 60")
        doc_calha.add_paragraph(f"Resultado Final: Q_proj = {q_proj1:.2f} L/min").bold = True
        
        doc_calha.add_heading(f'2. CAPACIDADE DA CALHA ({tipo_calha.upper()})', level=1)
        doc_calha.add_paragraph(f"Área Molhada (S) = {S_area:.4f} m²\nPerímetro Molhado (P) = {P_perimetro:.4f} m\nRaio Hidráulico (Rh) = {R_hidr:.4f} m\nRugosidade (n) = {n_rug:.3f}\nDeclividade (i) = {declividade:.4f}\nConstante (K) = 60000")
        formatar_formula_word(doc_calha, "Fórmula Base: ", "Q_calha = (K * S * Rh^(2/3) * i^(0.5)) / n")
        formatar_formula_word(doc_calha, "Substituição: ", f"Q_calha = (60000 * {S_area:.4f} * {R_hidr:.4f}^(2/3) * {declividade:.4f}^(0.5)) / {n_rug:.3f}")
        doc_calha.add_paragraph(f"Resultado Final: Q_calha = {q_calha_calc:.2f} L/min").bold = True
        
        status_txt = "ATENDE" if q_calha_calc >= q_proj1 else "NÃO ATENDE"
        doc_calha.add_heading(f'STATUS: {status_txt}', level=2)

        bio_calha = io.BytesIO()
        doc_calha.save(bio_calha)
        st.download_button(
            label="📄 Baixar Memorial em Word (.docx)", 
            data=bio_calha.getvalue(), 
            file_name="Memorial_Calhas.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ==============================================================================
# ABA 2: CONDUTORES VERTICAIS
# ==============================================================================
with tab2:
    st.header("Dimensionamento de Condutores Verticais")
    col_v1, col_v2 = st.columns([1, 1.2])
    
    with col_v1:
        st.subheader("Parâmetros de Entrada")
        area_c2 = st.number_input("Área de contribuição para este condutor (A) [m²]:", min_value=0.0, value=150.0, step=10.0, key="area_vert")
        i_pluv2 = st.number_input("Intensidade pluviométrica (I) [mm/h]:", min_value=0.0, value=150.0, step=10.0, key="i_vert")
        q_proj2 = (i_pluv2 * area_c2) / 60.0
        
        st.markdown("---")
        diametro_escolhido = st.selectbox(
            "Após realizar a leitura livre no ábaco ao lado, selecione o Diâmetro Nominal (DN) adotado:",
            ["70 mm (Mínimo)", "75 mm", "100 mm", "125 mm", "150 mm", "200 mm"]
        )

    with col_v2:
        st.subheader("📋 MEMORIAL DE CÁLCULO E ÁBACOS")
        
        st.markdown("### 1. Vazão de Projeto ($Q_{proj}$)")
        with st.container():
            st.markdown("**💡 1. Símbolos e Valores Cadastrados:**")
            st.markdown(f"* $I$ = **{i_pluv2:.2f}** mm/h\n* $A$ = **{area_c2:.2f}** m²")
            st.markdown("**📖 2. Fórmula Base:**")
            st.latex(r"Q_{proj} = \frac{I \cdot A}{60}")
            st.markdown("**🧪 3. Substituição Numérica:**")
            st.latex(f"Q_{{proj}} = \\frac{{ {i_pluv2:.2f} \\cdot {area_c2:.2f} }}{{ 60 }}")
            st.metric(label="Vazão a Escoar Calculada", value=f"{q_proj2:.2f} L/min")
        
        st.markdown("---")
        st.markdown("### 2. Ábaco Interativo de Verificação (Figura 3)")
        
        tipo_abaco = st.radio("Selecione o tipo de saída para o gráfico:", ["Saída em Aresta Viva", "Saída com Funil"])
        if tipo_abaco == "Saída em Aresta Viva":
            st.markdown("**Figura 3 (a) - Saída em Aresta Viva**")
            arquivo_img = "Abaco calha com saida em aresta viva.png"
        else:
            st.markdown("**Figura 3 (b) - Saída com Funil**")
            arquivo_img = "Abaco calha com funil da saida.png"
            
        st.info("🖱️ **Dica de Uso:** Clique e **arraste as linhas vermelha e azul** com o mouse diretamente sobre o gráfico para fazer sua leitura de forma dinâmica!")
        
        fig = go.Figure()
        if os.path.exists(arquivo_img):
            img = carregar_imagem(arquivo_img)
            fig.add_layout_image(
                dict(source=img, xref="x", yref="y", x=0, y=100, sizex=100, sizey=100, sizing="stretch", opacity=1, layer="below")
            )
        else:
            st.warning(f"⚠️ Imagem '{arquivo_img}' não encontrada no diretório do projeto. O gráfico aparecerá sem fundo.")
            
        fig.update_xaxes(range=[0, 100], showgrid=False, zeroline=False, visible=False)
        fig.update_yaxes(range=[0, 100], showgrid=False, zeroline=False, visible=False)
        
        # LINHAS ARRASTÁVEIS: 'editable=True' ativado para interação via JavaScript
        fig.add_shape(type="line", x0=25, y0=0, x1=25, y1=100, 
                      line=dict(color="red", width=2, dash="dash"), editable=True)
                      
        fig.add_shape(type="line", x0=0, y0=40, x1=100, y1=40, 
                      line=dict(color="blue", width=2), editable=True)
                      
        fig.update_layout(height=650, margin=dict(l=0, r=0, t=10, b=0), template="plotly_white")
        
        # Configuração para permitir a edição de formas no Plotly
        config_interativa = {'edits': {'shapePosition': True}}
        st.plotly_chart(fig, use_container_width=True, config=config_interativa)

        st.markdown(f"**Resultado Verificado no Ábaco:** Diâmetro Adotado = **{diametro_escolhido}**")

        # GERAÇÃO DO ARQUIVO WORD
        doc_vert = Document()
        criar_cabecalho_word(doc_vert, 'Memorial de Cálculo - Condutores Verticais')
        doc_vert.add_heading('1. VAZÃO DE PROJETO (Q_proj)', level=1)
        doc_vert.add_paragraph(f"Intensidade (I) = {i_pluv2:.2f} mm/h\nÁrea (A) = {area_c2:.2f} m²")
        formatar_formula_word(doc_vert, "Fórmula Base: ", "Q_proj = (I * A) / 60")
        formatar_formula_word(doc_vert, "Substituição: ", f"Q_proj = ({i_pluv2:.2f} * {area_c2:.2f}) / 60")
        doc_vert.add_paragraph(f"Resultado Final: Q_proj = {q_proj2:.2f} L/min").bold = True
        
        doc_vert.add_heading('2. LEITURA PELO ÁBACO (FIGURA 3)', level=1)
        doc_vert.add_paragraph(f"Tipo de Saída Analisada: {tipo_abaco}")
        doc_vert.add_paragraph("Análise realizada visualmente através do cruzamento de coordenadas do ábaco.")
        doc_vert.add_paragraph(f"RESULTADO DE PROJETO ADOTADO: {diametro_escolhido}").bold = True

        bio_vert = io.BytesIO()
        doc_vert.save(bio_vert)
        st.download_button(
            label="📄 Baixar Memorial em Word (.docx)", 
            data=bio_vert.getvalue(), 
            file_name="Memorial_Verticais.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ==============================================================================
# ABA 3: CONDUTORES HORIZONTAIS
# ==============================================================================
with tab3:
    st.header("Dimensionamento de Condutores Horizontais")
    col_h1, col_h2 = st.columns([1, 1.2])
    
    with col_h1:
        st.subheader("Parâmetros de Entrada")
        area_c3 = st.number_input("Área de contribuição (A) [m²]:", min_value=0.0, value=150.0, step=10.0, key="area_horiz")
        i_pluv3 = st.number_input("Intensidade pluviométrica (I) [mm/h]:", min_value=0.0, value=150.0, step=10.0, key="i_horiz")
        
        st.markdown("---")
        st.markdown("**Parâmetros do Trecho Horizontal**")
        opcoes_mat = {
            "Plástico, fibrocimento, aço galvanizado (n = 0,011)": 0.011,
            "Ferro fundido, concreto alisado (n = 0,012)": 0.012,
            "Cerâmica, concreto não alisado (n = 0,013)": 0.013
        }
        escolha_material = st.selectbox("Selecione o material para aplicar o fator de correção:", list(opcoes_mat.keys()))
        n_escolhido = opcoes_mat[escolha_material]
        fator_correcao = 0.011 / n_escolhido
        
        declividade_h = st.selectbox("Declividade do tubo (i):", ["0,5%", "1%", "2%", "4%"])
        diametro_tubo = st.selectbox("Diâmetro Nominal (DN) do tubo [mm]:", [50, 75, 100, 125, 150, 200, 250, 300], index=1)
        num_tubos = st.number_input("Quantidade de tubos paralelos:", min_value=1, value=1, step=1)

    with col_h2:
        st.subheader("📋 MEMORIAL DE CÁLCULO E TABELA DE REFERÊNCIA")
        
        st.markdown("### 1. Vazão de Projeto ($Q_{proj}$)")
        q_proj3 = (i_pluv3 * area_c3) / 60.0
        
        with st.container():
            st.markdown("**💡 1. Símbolos e Valores Cadastrados:**")
            st.markdown(f"* $I$ = **{i_pluv3:.2f}** mm/h\n* $A$ = **{area_c3:.2f}** m²")
            st.markdown("**📖 2. Fórmula Base:**")
            st.latex(r"Q_{proj} = \frac{I \cdot A}{60}")
            st.markdown("**🧪 3. Substituição Numérica:**")
            st.latex(f"Q_{{proj}} = \\frac{{ {i_pluv3:.2f} \\cdot {area_c3:.2f} }}{{ 60 }}")
            st.metric(label="Vazão Crítica de Projeto", value=f"{q_proj3:.2f} L/min")
        
        st.markdown("---")
        st.markdown(f"### 2. Capacidade Hidráulica Corrigida (Item 5.7.2.2)")
        
        dados_tabela4_base = {
            "Diâmetro (mm)": [50, 75, 100, 125, 150, 200, 250, 300],
            "0,5%": [32.0, 95.0, 204.0, 370.0, 602.0, 1310.0, 2390.0, 3890.0],
            "1%": [45.0, 133.0, 287.0, 521.0, 847.0, 1840.0, 3370.0, 5490.0],
            "2%": [64.0, 188.0, 405.0, 735.0, 1190.0, 2600.0, 4770.0, 7760.0],
            "4%": [90.0, 267.0, 575.0, 1040.0, 1680.0, 3680.0, 6740.0, 10900.0]
        }
        df_tab4 = pd.DataFrame(dados_tabela4_base)
        df_exibicao = df_tab4.copy()
        for col in ["0,5%", "1%", "2%", "4%"]:
            df_exibicao[col] = (df_exibicao[col] * fator_correcao).round(2)
            
        st.dataframe(df_exibicao, hide_index=True)
        
        idx_tubo = df_tab4.index[df_tab4['Diâmetro (mm)'] == diametro_tubo].tolist()[0]
        q_base_tabela = df_tab4.at[idx_tubo, declividade_h]
        q_corrigida = q_base_tabela * fator_correcao
        capacidade_total = q_corrigida * num_tubos
        
        with st.container():
            st.markdown("**💡 1. Símbolos e Constantes Coletadas:**")
            st.markdown(f"* $Q_{{base}}$ = **{q_base_tabela:.2f} L/min**\n* $n_{{material}}$ = **{n_escolhido:.3f}**\n* $\\text{{Tubos}}$ = **{num_tubos}**")
            st.markdown("**📖 2. Fórmulas de Ajuste Normativo:**")
            st.latex(r"Q_{corrigida} = Q_{base} \cdot \left(\frac{0,011}{n_{material}}\right)")
            st.latex(r"Q_{total} = Q_{corrigida} \cdot \text{Tubos}")
            st.markdown("**🧪 3. Substituição e Multiplicação dos Fatores:**")
            st.latex(f"Q_{{corrigida}} = {q_base_tabela:.2f} \\cdot \\left( \\frac{{ 0,011 }}{{ {n_escolhido:.3f} }} \\right) = {q_corrigida:.2f} \\text{{ L/min}}")
            st.latex(f"Q_{{total}} = {q_corrigida:.2f} \\cdot {num_tubos}")
            st.metric(label="Capacidade Total Calculada (Q_total)", value=f"{capacidade_total:.2f} L/min")
        
        st.markdown("---")
        if capacidade_total >= q_proj3:
            st.success(f"✅ **CONFORMIDADE ATENDIDA:** A capacidade do sistema horizontal ({capacidade_total:.2f} L/min) suporta a vazão.")
        else:
            st.error(f"❌ **NÃO ATENDE:** A capacidade calculada ({capacidade_total:.2f} L/min) está abaixo da vazão requerida.")

        # GERAÇÃO DO ARQUIVO WORD
        doc_horiz = Document()
        criar_cabecalho_word(doc_horiz, 'Memorial de Cálculo - Condutores Horizontais')
        doc_horiz.add_heading('1. VAZÃO DE PROJETO (Q_proj)', level=1)
        doc_horiz.add_paragraph(f"Intensidade (I) = {i_pluv3:.2f} mm/h\nÁrea (A) = {area_c3:.2f} m²")
        formatar_formula_word(doc_horiz, "Fórmula Base: ", "Q_proj = (I * A) / 60")
        formatar_formula_word(doc_horiz, "Substituição: ", f"Q_proj = ({i_pluv3:.2f} * {area_c3:.2f}) / 60")
        doc_horiz.add_paragraph(f"Resultado Final: Q_proj = {q_proj3:.2f} L/min").bold = True
        
        doc_horiz.add_heading('2. CAPACIDADE DOS CONDUTORES (TABELA 4)', level=1)
        doc_horiz.add_paragraph(f"Diâmetro Adotado (DN) = {diametro_tubo} mm\nDeclividade (i) = {declividade_h}\nMaterial (n) = {n_escolhido:.3f}\nQ_base (Tabela 4) = {q_base_tabela:.2f} L/min\nQuantidade de tubos = {num_tubos}")
        formatar_formula_word(doc_horiz, "Fórmulas: ", "Q_corrigida = Q_base * (0.011 / n_material)\nQ_total = Q_corrigida * Tubos")
        formatar_formula_word(doc_horiz, "Substituição: ", f"Q_corrigida = {q_base_tabela:.2f} * (0.011 / {n_escolhido:.3f}) = {q_corrigida:.2f} L/min\nQ_total = {q_corrigida:.2f} * {num_tubos}")
        doc_horiz.add_paragraph(f"Resultado Final: Q_total = {capacidade_total:.2f} L/min").bold = True
        
        status_txt_h = "ATENDE" if capacidade_total >= q_proj3 else "NÃO ATENDE"
        doc_horiz.add_heading(f'STATUS: {status_txt_h}', level=2)

        bio_horiz = io.BytesIO()
        doc_horiz.save(bio_horiz)
        st.download_button(
            label="📄 Baixar Memorial em Word (.docx)", 
            data=bio_horiz.getvalue(), 
            file_name="Memorial_Horizontais.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )